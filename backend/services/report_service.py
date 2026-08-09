from pathlib import Path
from datetime import datetime
import json

from docx import Document


class ReportService:

    def __init__(self):
        self.report_dir = (
            Path(__file__).resolve().parent.parent
            / "generated_reports"
        )

        self.report_dir.mkdir(
            parents=True,
            exist_ok=True
        )

    def _parse_json(self, value, default=None):

        if default is None:
            default = []

        if not value:
            return default

        try:
            return json.loads(value)

        except (json.JSONDecodeError, TypeError):
            return default

    def generate_compliance_report(
        self,
        company,
        compliance_report,
        documents,
        recommendations,
        regulatory_alerts
    ):

        document = Document()

        # =================================================
        # TITLE
        # =================================================

        title = document.add_heading(
            "Compliance Assessment Report",
            level=0
        )

        title.alignment = 1

        document.add_paragraph(
            f"Generated: "
            f"{datetime.now().strftime('%d %B %Y, %H:%M')}"
        )

        # =================================================
        # COMPANY DETAILS
        # =================================================

        document.add_heading(
            "1. Company Details",
            level=1
        )

        company_table = document.add_table(
            rows=0,
            cols=2
        )

        company_details = [
            ("Company Name", company.company_name),
            ("CIN", company.cin),
            ("PAN", company.pan),
            ("GSTIN", company.gstin),
            ("Business Type", company.business_type),
            ("Industry", company.industry),
            ("State", company.state),
        ]

        for label, value in company_details:

            row = company_table.add_row().cells

            row[0].text = label
            row[1].text = str(
                value if value else "Not available"
            )

        # =================================================
        # COMPLIANCE SUMMARY
        # =================================================

        document.add_heading(
            "2. Compliance Summary",
            level=1
        )

        if compliance_report:

            document.add_paragraph(
                f"Compliance Score: "
                f"{compliance_report.compliance_score}%"
            )

            document.add_paragraph(
                f"Risk Level: "
                f"{compliance_report.risk_level}"
            )

            present_documents = self._parse_json(
                compliance_report.present_documents
            )

            missing_documents = self._parse_json(
                compliance_report.missing_documents
            )

            document.add_heading(
                "Present Documents",
                level=2
            )

            if present_documents:

                for item in present_documents:

                    document.add_paragraph(
                        str(item),
                        style="List Bullet"
                    )

            else:

                document.add_paragraph(
                    "None"
                )

            document.add_heading(
                "Missing Documents",
                level=2
            )

            if missing_documents:

                for item in missing_documents:

                    document.add_paragraph(
                        str(item),
                        style="List Bullet"
                    )

            else:

                document.add_paragraph(
                    "None"
                )

            findings = self._parse_json(
                compliance_report.findings
            )

            document.add_heading(
                "Compliance Findings",
                level=2
            )

            if findings:

                for finding in findings:

                    document.add_paragraph(
                        f"{finding.get('document', 'Unknown')}: "
                        f"{finding.get('status', 'Unknown')}"
                    )

                    if finding.get("reason"):

                        document.add_paragraph(
                            f"Reason: "
                            f"{finding['reason']}"
                        )

            else:

                document.add_paragraph(
                    "No findings available."
                )

        else:

            document.add_paragraph(
                "No compliance assessment is available."
            )

        # =================================================
        # UPLOADED DOCUMENTS
        # =================================================

        document.add_heading(
            "3. Uploaded Documents",
            level=1
        )

        if documents:

            for item in documents:

                document.add_paragraph(
                    f"{item.document_type or 'Unknown'} — "
                    f"{item.filename} "
                    f"({item.status})",
                    style="List Bullet"
                )

        else:

            document.add_paragraph(
                "No documents have been uploaded."
            )

        # =================================================
        # RECOMMENDATIONS
        # =================================================

        document.add_heading(
            "4. Recommendations",
            level=1
        )

        if recommendations:

            for recommendation in recommendations:

                document.add_heading(
                    (
                        f"{recommendation.document} — "
                        f"{recommendation.priority}"
                    ),
                    level=2
                )

                document.add_paragraph(
                    f"Action: "
                    f"{recommendation.action}"
                )

                if recommendation.reason:

                    document.add_paragraph(
                        f"Reason: "
                        f"{recommendation.reason}"
                    )

                if recommendation.next_step:

                    document.add_paragraph(
                        f"Next Step: "
                        f"{recommendation.next_step}"
                    )

        else:

            document.add_paragraph(
                "No recommendations are currently available."
            )

        # =================================================
        # REGULATORY INTELLIGENCE ALERTS
        # =================================================

        document.add_heading(
            "5. Regulatory Intelligence Alerts",
            level=1
        )

        if regulatory_alerts:

            for alert in regulatory_alerts:

                document.add_heading(
                    (
                        f"{alert.title} — "
                        f"{alert.severity}"
                    ),
                    level=2
                )

                document.add_paragraph(
                    f"Authority: {alert.authority}"
                )

                document.add_paragraph(
                    f"Affected Document: "
                    f"{alert.affected_document}"
                )

                document.add_paragraph(
                    f"Status: {alert.status}"
                )

                document.add_paragraph(
                    f"Update ID: {alert.update_id}"
                )

        else:

            document.add_paragraph(
                "No regulatory alerts are currently recorded."
            )

        # =================================================
        # REPORT NOTE
        # =================================================

        document.add_heading(
            "6. Report Note",
            level=1
        )

        document.add_paragraph(
            "This report is generated from the compliance "
            "information and regulatory intelligence available "
            "in the prototype system at the time of generation."
        )

        # =================================================
        # PROFESSIONAL FILE NAME
        # =================================================

        safe_company_name = "".join(
            character
            if character.isalnum()
            else "_"
            for character in company.company_name
        ).strip("_")

        report_timestamp = datetime.now().strftime(
            "%Y-%m-%d_%H-%M"
        )

        filename = (
            f"{safe_company_name}_"
            f"Compliance_Report_"
            f"{report_timestamp}.docx"
        )

        report_path = (
            self.report_dir / filename
        )

        # =================================================
        # SAVE REPORT
        # =================================================

        document.save(report_path)

        # =================================================
        # RETURN PATH TO API
        # =================================================

        return report_path